"""
Document Processor
Handles PDF, DOCX, XLSX, and text files with proper table extraction
"""
import os
from io import BytesIO
from datetime import datetime
import json

# PDF processing
try:
    from PyPDF2 import PdfReader
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# Excel processing
try:
    import pandas as pd
    import openpyxl
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False
    pd = None
    openpyxl = None

# Word document processing
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    DocxDocument = None


class DocumentProcessor:
    """Process various document types and extract content"""
    
    def __init__(self):
        self.supported_types = {
            'pdf': ['.pdf'],
            'excel': ['.xlsx', '.xls', '.xlsm'],
            'word': ['.docx'],
            'text': ['.txt', '.md'],
            'html': ['.html', '.htm']
        }
    
    def detect_file_type(self, filename):
        """Detect file type from extension"""
        ext = os.path.splitext(filename.lower())[1]
        
        for file_type, extensions in self.supported_types.items():
            if ext in extensions:
                return file_type
        
        return 'unknown'
    
    def process_file(self, file_path_or_bytes, filename):
        """
        Process any supported file type
        
        Args:
            file_path_or_bytes: Either file path (str) or file bytes (BytesIO)
            filename: Original filename
        
        Returns:
            Dict with extracted content
        """
        file_type = self.detect_file_type(filename)
        
        processors = {
            'pdf': self.process_pdf,
            'excel': self.process_excel,
            'word': self.process_word,
            'text': self.process_text,
            'html': self.process_html
        }
        
        processor = processors.get(file_type)
        if processor:
            return processor(file_path_or_bytes, filename)
        else:
            return {
                'success': False,
                'error': f'Unsupported file type: {file_type}',
                'file_type': file_type
            }
    
    def process_pdf(self, file_path_or_bytes, filename):
        """Extract text and tables from PDF"""
        if not PDF_AVAILABLE:
            return {'success': False, 'error': 'PyPDF2 not installed'}
        
        try:
            # Read PDF
            if isinstance(file_path_or_bytes, str):
                reader = PdfReader(file_path_or_bytes)
            else:
                reader = PdfReader(file_path_or_bytes)
            
            # Extract text from all pages
            text_content = []
            for page_num, page in enumerate(reader.pages, 1):
                page_text = page.extract_text()
                if page_text.strip():
                    text_content.append({
                        'page': page_num,
                        'text': page_text
                    })
            
            # Combine all text
            full_text = '\n\n'.join([p['text'] for p in text_content])
            
            return {
                'success': True,
                'file_type': 'pdf',
                'filename': filename,
                'text': full_text,
                'pages': text_content,
                'page_count': len(reader.pages),
                'metadata': {
                    'page_count': len(reader.pages),
                    'has_text': len(full_text) > 100
                }
            }
        except Exception as e:
            return {
                'success': False,
                'error': f'PDF processing error: {str(e)}',
                'file_type': 'pdf'
            }
    
    def process_excel(self, file_path_or_bytes, filename):
        """Extract data from Excel files"""
        if not EXCEL_AVAILABLE:
            return {'success': False, 'error': 'pandas/openpyxl not installed'}
        
        try:
            # Read Excel file
            if isinstance(file_path_or_bytes, str):
                xls = pd.ExcelFile(file_path_or_bytes)
            else:
                xls = pd.ExcelFile(file_path_or_bytes)
            
            # Process each sheet
            sheets_data = []
            full_text = []
            
            for sheet_name in xls.sheet_names:
                df = pd.read_excel(xls, sheet_name=sheet_name)
                
                # Convert to text summary
                sheet_text = f"Sheet: {sheet_name}\n"
                sheet_text += f"Columns: {', '.join(df.columns.tolist())}\n"
                sheet_text += f"Rows: {len(df)}\n\n"
                
                # Add data summary
                for idx, row in df.head(10).iterrows():  # First 10 rows for text
                    row_text = ' | '.join([f"{col}: {val}" for col, val in row.items()])
                    sheet_text += f"{row_text}\n"
                
                if len(df) > 10:
                    sheet_text += f"... and {len(df) - 10} more rows\n"
                
                sheets_data.append({
                    'sheet_name': sheet_name,
                    'columns': df.columns.tolist(),
                    'row_count': len(df),
                    'data': df.to_dict('records'),  # Full structured data
                    'text_summary': sheet_text
                })
                
                full_text.append(sheet_text)
            
            return {
                'success': True,
                'file_type': 'excel',
                'filename': filename,
                'text': '\n\n'.join(full_text),
                'sheets': sheets_data,
                'sheet_count': len(xls.sheet_names),
                'metadata': {
                    'sheet_names': xls.sheet_names,
                    'total_sheets': len(xls.sheet_names)
                }
            }
        except Exception as e:
            return {
                'success': False,
                'error': f'Excel processing error: {str(e)}',
                'file_type': 'excel'
            }
    
    def process_word(self, file_path_or_bytes, filename):
        """Extract text from Word documents"""
        if not DOCX_AVAILABLE:
            return {'success': False, 'error': 'python-docx not installed'}
        
        try:
            # Read DOCX
            if isinstance(file_path_or_bytes, str):
                doc = DocxDocument(file_path_or_bytes)
            else:
                doc = DocxDocument(file_path_or_bytes)
            
            # Extract paragraphs
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            full_text = '\n\n'.join(paragraphs)
            
            # Extract tables
            tables_data = []
            for table_idx, table in enumerate(doc.tables):
                table_text = []
                for row in table.rows:
                    row_text = ' | '.join([cell.text for cell in row.cells])
                    table_text.append(row_text)
                
                tables_data.append({
                    'table_index': table_idx,
                    'rows': len(table.rows),
                    'columns': len(table.columns),
                    'text': '\n'.join(table_text)
                })
            
            # Add tables to full text
            if tables_data:
                full_text += '\n\nTables:\n'
                for table in tables_data:
                    full_text += f"\nTable {table['table_index']}:\n{table['text']}\n"
            
            return {
                'success': True,
                'file_type': 'word',
                'filename': filename,
                'text': full_text,
                'paragraphs': paragraphs,
                'tables': tables_data,
                'metadata': {
                    'paragraph_count': len(paragraphs),
                    'table_count': len(tables_data)
                }
            }
        except Exception as e:
            return {
                'success': False,
                'error': f'Word processing error: {str(e)}',
                'file_type': 'word'
            }
    
    def process_text(self, file_path_or_bytes, filename):
        """Extract text from plain text files"""
        try:
            if isinstance(file_path_or_bytes, str):
                with open(file_path_or_bytes, 'r', encoding='utf-8') as f:
                    text = f.read()
            else:
                text = file_path_or_bytes.read().decode('utf-8')
            
            return {
                'success': True,
                'file_type': 'text',
                'filename': filename,
                'text': text,
                'metadata': {
                    'char_count': len(text),
                    'line_count': len(text.split('\n'))
                }
            }
        except Exception as e:
            return {
                'success': False,
                'error': f'Text processing error: {str(e)}',
                'file_type': 'text'
            }
    
    def process_html(self, file_path_or_bytes, filename):
        """Extract text from HTML files"""
        try:
            from bs4 import BeautifulSoup
            
            if isinstance(file_path_or_bytes, str):
                with open(file_path_or_bytes, 'r', encoding='utf-8') as f:
                    html = f.read()
            else:
                html = file_path_or_bytes.read().decode('utf-8')
            
            soup = BeautifulSoup(html, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Get text
            text = soup.get_text()
            
            # Clean up text
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = '\n'.join(chunk for chunk in chunks if chunk)
            
            return {
                'success': True,
                'file_type': 'html',
                'filename': filename,
                'text': text,
                'metadata': {
                    'char_count': len(text)
                }
            }
        except Exception as e:
            return {
                'success': False,
                'error': f'HTML processing error: {str(e)}',
                'file_type': 'html'
            }


# Global instance
_processor = None

def get_processor():
    """Get or create global processor instance"""
    global _processor
    if _processor is None:
        _processor = DocumentProcessor()
    return _processor


if __name__ == "__main__":
    # Test document processor
    print("Testing Document Processor...")
    
    processor = DocumentProcessor()
    
    # Test text processing
    test_text = BytesIO(b"This is a test document.\nIt has multiple lines.")
    result = processor.process_file(test_text, "test.txt")
    
    if result['success']:
        print(f"✅ Processed {result['filename']}")
        print(f"   Type: {result['file_type']}")
        print(f"   Text length: {len(result['text'])} chars")
    else:
        print(f"❌ Error: {result['error']}")
